# -*- Product under GNU GPL v3 -*-
# -*- Author: E.Aivayan -*-
from pathlib import Path
from unittest.mock import create_autospec, MagicMock
from shutil import copyfile

from docx import Document

from eaiBat.evidenceGenerators.DocxEvidence import _file_to_evidence


class TestFileToEvidence:
    def test_valide_img_file(self, tmp_path):
        docu = create_autospec(Document, instance=True)
        docu.add_paragraph = MagicMock()
        docu.add_picture = MagicMock()
        docu.add_page_break = MagicMock()
        picture = Path("tests/resources/test_img.png")
        copyfile(picture.absolute(), tmp_path / "test_img.png")
        Path.mkdir(tmp_path / "included_files", exist_ok=True)
        _file_to_evidence(docu,
                          (str(tmp_path / "test_img.png"), "img"),
                          tmp_path / "included_files",
                          str(tmp_path.absolute()))
        assert (tmp_path / "included_files" / "test_img.png").exists(), "Moved file not found"
        assert (tmp_path / "included_files" / "test_img.png").is_file(), "Moved file is not file"

    def test_valid_other_file(self, tmp_path):
        docu = create_autospec(Document, instance=True)
        docu.add_paragraph = MagicMock()
        docu.add_picture = MagicMock()
        docu.add_page_break = MagicMock()
        picture = Path("tests/resources/test_img.png")
        copyfile(picture.absolute(), tmp_path / "test_img.png")
        Path.mkdir(tmp_path / "included_files", exist_ok=True)
        _file_to_evidence(docu,
                          (str(tmp_path / "test_img.png"), "txt"),
                          tmp_path / "included_files",
                          str(tmp_path.absolute()))
        assert (tmp_path / "included_files" / "test_img.png").exists(), "Moved file not found"
        assert (tmp_path / "included_files" / "test_img.png").is_file(), "Moved file is not file"

    def test_invalid_path_file(self, tmp_path):
        docu = create_autospec(Document, instance=True)
        docu.add_paragraph = MagicMock()
        docu.add_picture = MagicMock()
        docu.add_page_break = MagicMock()
        picture = Path("tests/resources/test_img.png")
        copyfile(picture.absolute(), tmp_path / "test_img.png")
        Path.mkdir(tmp_path / "included_files", exist_ok=True)
        _file_to_evidence(docu,
                          (str(tmp_path / "C:/test_img.png"), "img"),
                          tmp_path / "included_files",
                          str(tmp_path.absolute()))
        assert (tmp_path / "included_files" / "test_img.png").exists(), "Moved file not found"
        assert (tmp_path / "included_files" / "test_img.png").is_file(), "Moved file is not file"

    def test_invalid_file(self, tmp_path):
        docu = create_autospec(Document, instance=True)
        docu.add_paragraph = MagicMock()
        docu.add_picture = MagicMock()
        docu.add_page_break = MagicMock()
        picture = Path("tests/resources/test_img.png")
        copyfile(picture.absolute(), tmp_path / "test_img.png")
        Path.mkdir(tmp_path / "included_files", exist_ok=True)
        _file_to_evidence(docu,
                          (str(tmp_path / "C:/test_img/"), "img"),
                          tmp_path / "included_files",
                          str(tmp_path.absolute()))
        assert not (tmp_path / "included_files" / "test_img.png").exists(), "Moved file found"
