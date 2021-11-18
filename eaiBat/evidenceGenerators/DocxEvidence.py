# -*- Product under GNU GPL v3 -*-
# -*- Author: E.Aivayan -*-
import os
import json

from docx import Document
from docx.shared import Cm
from logging import getLogger
from pathlib import Path
from shutil import move
from xml.dom import minidom
from requests.models import Response

log = getLogger(__name__)


def generate_docx_evidence(evidence_folder, evidence_filename, history):
    evidence_document = Document()
    evidence_document.add_heading(evidence_filename, 0)
    included_file_folder = Path(evidence_folder) / "included_files"
    Path.mkdir(included_file_folder, exist_ok=True)
    for step, events in history.items():
        evidence_document.add_heading(f"{step[0]}: {step[1]}", 2)
        for event in events:
            try:
                if isinstance(event, str):
                    evidence_document.add_paragraph(event)
                elif isinstance(event, tuple):
                    _file_to_evidence(evidence_document,
                                      event,
                                      included_file_folder,
                                      evidence_folder)
                elif isinstance(event, dict):
                    _dict_to_evidence(evidence_document,
                                      event,
                                      included_file_folder,
                                      evidence_folder)
                else:
                    _response_to_evidence(evidence_document,
                                          event)
            except Exception as exception:
                log.error(exception)

    file_path = Path(evidence_folder)
    save_to = file_path / evidence_filename
    log.debug(f"Save document to '{save_to.absolute()}'")
    evidence_document.save(save_to.absolute())


def _file_to_evidence(docx_document: Document, event: tuple, included_files_folder: Path,
                      evidence_folder: str):
    log.debug("Try to include external file into the document")
    path_from_event = Path(event[0])
    # First check the event file element is pointing to a file
    if path_from_event.exists() and path_from_event.is_file():
        log.debug("File found")
        file_path = path_from_event
    else:  # The event file does not exist or is not pointing to a file
        log.debug("File not found")
        # Try to create a filepath with parent destination folder and filename
        file_path = Path(evidence_folder) / path_from_event.name

    log.debug(f"Selected path is '{file_path}'")
    try:
        if file_path.exists() and file_path.is_file():
            log.debug(f"Move the file from {file_path} to {included_files_folder}/{file_path.name}")
            move(file_path, included_files_folder / file_path.name)
            new_path = included_files_folder / file_path.name
            if event[1].casefold() == "img":
                log.debug(f"Try to include '{new_path}'")
                log.debug(f"Add this moved picture '{new_path.absolute()}'")
                docx_document.add_picture(str(new_path.absolute()), width=Cm(18))
                docx_document.add_page_break()
                log.debug("Picture included")
            else:
                docx_document.add_paragraph(
                    f"A file has been produced or retrieved within this step."
                    f"You could see it there : '{new_path}'")
        else:
            docx_document.add_paragraph(f"{event[1]} file is located at {event[0]} (but not found)")
    except Exception as exception:
        log.error(exception.args)
        docx_document.add_paragraph(f"{event[1]} file is located at {event[0]} (but not found)")


def _dict_to_evidence(docx_document: Document, event: dict, file, storage):
    for key, value in event.items():
        docx_document.add_heading(f"{key}", 3)
        if isinstance(value, Response):
            _response_to_evidence(docx_document, value, sub_level=True)
        elif isinstance(value, tuple):
            _file_to_evidence(docx_document, value, file, storage)
        else:
            docx_document.add_paragraph(value)


def _response_to_evidence(docx_document: Document,
                          event: Response,
                          sub_level: bool = False):
    section_number = 4 if sub_level else 3

    docx_document.add_heading("URL", section_number)
    docx_document.add_paragraph(str(event.url))
    docx_document.add_heading("VERB", section_number)
    docx_document.add_paragraph(str(event.request.method))
    docx_document.add_heading("Request headers", section_number)
    docx_document.add_paragraph(json.dumps(dict(event.request.headers), indent=2))

    if event.request.body is not None:
        docx_document.add_heading("Body", section_number)
        docx_document.add_paragraph(json.dumps(json.loads(event.request.body), indent=2))

    docx_document.add_heading("Response status code", section_number)
    docx_document.add_paragraph(str(event.status_code))
    docx_document.add_heading("Response headers", section_number)
    docx_document.add_paragraph(json.dumps(dict(event.headers), indent=2))
    docx_document.add_heading("Response content")
    try:
        response = json.dumps(event.json(), indent=2)
        docx_document.add_paragraph("Json response")
        docx_document.add_paragraph(response)
        return
    except Exception as exception:
        log.info(f"Response is not a json.\n Get {exception.args[0]}")

    try:
        response = minidom.parseString(event.text).toprettyxml(indent="   ")
        docx_document.add_paragraph("XML response")
        docx_document.add_paragraph(response)
        return
    except Exception as exception:
        log.info(f"Response is not a XML.\n Get {exception.args[0]}")

    docx_document.add_paragraph("Raw response")
    docx_document.add_paragraph(event.text)
